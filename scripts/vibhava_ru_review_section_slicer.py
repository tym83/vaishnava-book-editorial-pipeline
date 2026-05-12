#!/usr/bin/env python3
# Copyright 2026 The Vaishnava Book Editorial Pipeline Authors
# SPDX-License-Identifier: Apache-2.0

"""Build the RU Vibhava Tom 1 review split.

The current style-based split has 81 files because the final three chapter
headings are styled as body text inside 081.docx.  For source/translation review
we need a section split matching the English PDF: 084 sections total.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import List, Sequence

from docx import Document

from chapter_splitter import (
    Block,
    copy_docx_chapter,
    docx_blocks,
    normalize_text,
)


TAIL_SECTIONS = [
    ("081", "Чатур-масья и другие обеты"),
    ("082", "Здоровье"),
    ("083", "Другие наставления и истории"),
    ("084", "Его вечная форма и внутренний экстаз"),
]


@dataclass
class SectionInfo:
    code: str
    title: str
    source: str
    output: str
    start_idx: int
    end_idx: int
    words: int


def die(message: str) -> None:
    raise SystemExit(f"ERROR: {message}")


def non_section_prefix(blocks: Sequence[Block], idx: int) -> int:
    """Exclude preceding chapter-number paragraphs from section ranges."""

    previous = None
    for block in reversed(blocks):
        if block.index >= idx:
            continue
        if block.kind == "p" and block.text.strip():
            previous = block
            break
    if previous and re.fullmatch(r"Глава\s+\d+", previous.text.strip(), flags=re.IGNORECASE):
        return previous.index
    return idx


def find_title_block(blocks: Sequence[Block], title: str) -> Block:
    target = normalize_text(title)
    matches = [
        block
        for block in blocks
        if block.kind == "p" and block.norm_text == target
    ]
    if not matches:
        die(f"Could not find tail title in 081.docx: {title}")
    if len(matches) > 1:
        die(f"Ambiguous tail title in 081.docx: {title}")
    return matches[0]


def range_word_count(blocks: Sequence[Block], start_idx: int, end_idx: int) -> int:
    return sum(
        len(block.text.split())
        for block in blocks
        if start_idx <= block.index < end_idx and block.text.strip()
    )


def copy_head_sections(source_dir: Path, output_dir: Path) -> List[SectionInfo]:
    infos: List[SectionInfo] = []
    for code in [f"{i:03d}" for i in range(1, 81)]:
        src = source_dir / f"{code}.docx"
        if not src.exists():
            die(f"Missing source section: {src}")
        dst = output_dir / f"{code}.docx"
        shutil.copy2(src, dst)
        doc = Document(str(dst))
        words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
        title = next((p.text.strip() for p in doc.paragraphs if p.text.strip()), "")
        infos.append(
            SectionInfo(
                code=code,
                title=title,
                source=str(src),
                output=str(dst),
                start_idx=0,
                end_idx=0,
                words=words,
            )
        )
    return infos


def split_tail(source_dir: Path, output_dir: Path) -> List[SectionInfo]:
    src = source_dir / "081.docx"
    if not src.exists():
        die(f"Missing tail source section: {src}")
    doc = Document(str(src))
    blocks = docx_blocks(doc)
    max_idx = max(block.index for block in blocks if block.kind != "sectPr") + 1

    starts = []
    for code, title in TAIL_SECTIONS:
        block = find_title_block(blocks, title)
        starts.append((code, title, block.index))

    infos: List[SectionInfo] = []
    for i, (code, title, start_idx) in enumerate(starts):
        if i + 1 < len(starts):
            end_idx = non_section_prefix(blocks, starts[i + 1][2])
        else:
            end_idx = max_idx
        dst = output_dir / f"{code}.docx"
        copy_docx_chapter(src, blocks, start_idx, end_idx, dst)
        infos.append(
            SectionInfo(
                code=code,
                title=title,
                source=str(src),
                output=str(dst),
                start_idx=start_idx,
                end_idx=end_idx,
                words=range_word_count(blocks, start_idx, end_idx),
            )
        )
    return infos


def write_index(output_dir: Path, infos: Sequence[SectionInfo]) -> None:
    payload = {
        "output_dir": str(output_dir),
        "sections": [asdict(info) for info in infos],
    }
    (output_dir / "index.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# Vibhava Tom 1 RU Review Section Index",
        "",
        "| Code | Title | Words |",
        "|---|---|---:|",
    ]
    for info in infos:
        lines.append(f"| {info.code} | {info.title} | {info.words} |")
    (output_dir / "index.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_review_split(source_dir: Path, output_dir: Path) -> List[SectionInfo]:
    output_dir.mkdir(parents=True, exist_ok=True)
    infos = copy_head_sections(source_dir, output_dir)
    infos.extend(split_tail(source_dir, output_dir))
    write_index(output_dir, infos)
    return infos


def main() -> int:
    parser = argparse.ArgumentParser(description="Build 84-section RU split for Vibhava Tom 1 review")
    parser.add_argument("source_dir", help="Existing style-based split directory containing 001.docx..081.docx")
    parser.add_argument("--output-dir", required=True)
    args = parser.parse_args()

    infos = build_review_split(Path(args.source_dir), Path(args.output_dir))
    print(f"Wrote {len(infos)} sections to {args.output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
