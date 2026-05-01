#!/usr/bin/env python3
# Copyright 2026 The Vaishnava Book Editorial Pipeline Authors
# SPDX-License-Identifier: Apache-2.0

"""
Split books into chapter files.

v1 scope:
- docx/doc -> split into docx while preserving document styles/formatting
- pdf -> split into chapter PDF files while preserving original page geometry/layout
- chapter detection:
  * by paragraph style names
  * by explicit title list
  * by TOC extraction heuristics

Examples:
  python3 chapter_splitter.py extract-titles book.docx --mode toc --toc-marker "Содержание"
  python3 chapter_splitter.py split book.docx --mode toc --toc-marker "Содержание" --output-dir out
  python3 chapter_splitter.py split book.doc --mode style --heading-style "Заголовок 1" --output-dir out
  python3 chapter_splitter.py split book.pdf --mode toc --toc-marker "Contents" --output-dir out
"""

from __future__ import annotations

import argparse
import copy
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Sequence, Tuple

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn


DEFAULT_HEADING_STYLES = ["Заголовок 1", "Heading 1", "H1", "Chapter Title"]
DEFAULT_TOC_MARKERS = ["Содержание", "Contents", "Оглавление"]
PARA_TAG = qn("w:p")
TBL_TAG = qn("w:tbl")
SECTPR_TAG = qn("w:sectPr")
TEXT_TAG = qn("w:t")
PSTYLE_TAG = qn("w:pStyle")


@dataclass
class Block:
    index: int
    kind: str
    text: str
    norm_text: str
    style_name: Optional[str]
    xml: object


def die(msg: str) -> None:
    print(f"ERROR: {msg}", file=sys.stderr)
    raise SystemExit(1)


def normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text.strip())
    text = text.replace("ё", "е").replace("Ё", "Е")
    return text.casefold()


def visible_text_for_xml(node) -> str:
    return "".join(t.text for t in node.iter(TEXT_TAG) if t.text)


def style_id_to_name_map(doc: Document) -> dict:
    mapping = {}
    for style in doc.styles:
        style_id = getattr(style, "style_id", None)
        name = getattr(style, "name", None)
        if style_id and name:
            mapping[style_id] = name
    return mapping


def docx_blocks(doc: Document) -> List[Block]:
    styles = style_id_to_name_map(doc)
    blocks: List[Block] = []
    for idx, child in enumerate(doc.element.body.iterchildren()):
        if child.tag == SECTPR_TAG:
            blocks.append(Block(idx, "sectPr", "", "", None, child))
            continue
        if child.tag == PARA_TAG:
            text = visible_text_for_xml(child).strip()
            p_style = None
            ppr = child.find(qn("w:pPr"))
            if ppr is not None:
                pstyle = ppr.find(PSTYLE_TAG)
                if pstyle is not None:
                    sid = pstyle.get(qn("w:val"))
                    p_style = styles.get(sid, sid)
            blocks.append(Block(idx, "p", text, normalize_text(text), p_style, child))
            continue
        if child.tag == TBL_TAG:
            text = visible_text_for_xml(child).strip()
            blocks.append(Block(idx, "tbl", text, normalize_text(text), None, child))
            continue
        blocks.append(Block(idx, "other", "", "", None, child))
    return blocks


def convert_doc_to_docx(path: Path) -> Path:
    temp_dir = Path(tempfile.mkdtemp(prefix="chapter-splitter-"))
    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(temp_dir),
            str(path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    converted = temp_dir / f"{path.stem}.docx"
    if not converted.exists():
        die(f"Could not convert {path} to docx")
    return converted


def extract_pdf_text(path: Path) -> str:
    proc = subprocess.run(
        ["pdftotext", "-layout", str(path), "-"],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    return proc.stdout


def pdf_page_count(path: Path) -> int:
    proc = subprocess.run(
        ["pdfinfo", str(path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    match = re.search(r"^Pages:\s+(\d+)$", proc.stdout, flags=re.MULTILINE)
    if not match:
        die(f"Could not determine page count for {path}")
    return int(match.group(1))


def extract_pdf_page_text(path: Path, page: int) -> str:
    proc = subprocess.run(
        ["pdftotext", "-layout", "-f", str(page), "-l", str(page), str(path), "-"],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    return proc.stdout


def is_title_candidate(line: str, max_len: int = 120) -> bool:
    text = line.strip()
    if not text:
        return False
    if len(text) > max_len:
        return False
    if re.search(r"[.!?]$", text):
        return False
    if len(text.split()) > 14:
        return False
    return True


def merge_title_lines(lines: Sequence[str]) -> List[str]:
    merged: List[str] = []
    for raw in lines:
        line = re.sub(r"\s+", " ", raw.strip())
        if not line:
            continue
        if (
            merged
            and line[:1].islower()
            and len(merged[-1]) + 1 + len(line) <= 180
        ):
            merged[-1] = f"{merged[-1]} {line}"
        else:
            merged.append(line)
    return merged


def dedupe_preserve_order(lines: Sequence[str]) -> List[str]:
    seen = set()
    out = []
    for line in lines:
        norm = normalize_text(line)
        if norm in seen:
            continue
        seen.add(norm)
        out.append(line)
    return out


def clean_pdf_line(text: str) -> str:
    text = text.replace("\x0c", " ").replace("\x08", " ")
    text = text.replace("\ufffd", " ")
    text = re.sub(r"[\u200b\ufeff]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def strip_pdf_toc_entry(line: str) -> Optional[str]:
    text = clean_pdf_line(line)
    if not text:
        return None
    page_match = re.search(r"(?:^|\s)(\d+|[ivxlcdm]+)\s*$", text, flags=re.IGNORECASE)
    if not page_match:
        return None
    title = text[:page_match.start()].strip()
    title = re.sub(r"[.\-•·_]+$", "", title).strip()
    title = re.sub(r"\s+", " ", title)
    return title or None


def normalize_pdf_match_text(text: str) -> str:
    text = clean_pdf_line(text)
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[^\w\s/’'&-]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return normalize_text(text)


def title_matches_page_line(title: str, page_line: str) -> bool:
    title_norm = normalize_pdf_match_text(title)
    line_norm = normalize_pdf_match_text(page_line)
    if not title_norm or not line_norm:
        return False
    if title_norm == line_norm:
        return True
    if line_norm.startswith(title_norm) or title_norm.startswith(line_norm):
        return True
    title_tokens = title_norm.split()
    line_tokens = line_norm.split()
    if title_tokens and len(title_tokens) <= len(line_tokens):
        if line_tokens[: len(title_tokens)] == title_tokens:
            return True
    return False


def find_toc_start_index(blocks: Sequence[Block], toc_markers: Sequence[str]) -> Optional[int]:
    markers = {normalize_text(m) for m in toc_markers}
    for block in blocks:
        if block.kind == "p" and block.norm_text in markers:
            return block.index
    return None


def extract_titles_from_docx_toc(
    blocks: Sequence[Block],
    toc_markers: Sequence[str],
    max_len: int,
) -> Tuple[List[str], Optional[int], Optional[int]]:
    toc_start = find_toc_start_index(blocks, toc_markers)
    if toc_start is None:
        return [], None, None

    titles: List[str] = []
    collecting = False
    stop_index = None
    for block in blocks:
        if block.index <= toc_start or block.kind != "p":
            continue
        text = re.sub(r"\s+", " ", block.text.strip())
        if not collecting:
            if is_title_candidate(text, max_len=max_len):
                collecting = True
                titles.append(text)
                continue
            continue

        if is_title_candidate(text, max_len=max_len):
            titles.append(text)
            continue

        stop_index = block.index
        break

    titles = dedupe_preserve_order(merge_title_lines(titles))
    return titles, toc_start, stop_index


def extract_titles_from_docx_styles(blocks: Sequence[Block], heading_styles: Sequence[str]) -> List[str]:
    style_set = {s.casefold() for s in heading_styles}
    titles = [
        block.text.strip()
        for block in blocks
        if block.kind == "p" and block.style_name and block.style_name.casefold() in style_set and block.text.strip()
    ]
    return dedupe_preserve_order(titles)


def load_title_list(path: Path) -> List[str]:
    lines = path.read_text(encoding="utf-8").splitlines()
    return dedupe_preserve_order([line.strip() for line in lines if line.strip()])


def find_docx_split_points_from_titles(
    blocks: Sequence[Block],
    titles: Sequence[str],
    search_start_index: int = 0,
) -> List[Tuple[str, int]]:
    points: List[Tuple[str, int]] = []
    pos = search_start_index
    for title in titles:
        target = normalize_text(title)
        found = None
        for block in blocks:
            if block.index < pos or block.kind != "p":
                continue
            if block.norm_text == target:
                found = block.index
                break
        if found is None:
            continue
        points.append((title, found))
        pos = found + 1
    return points


def find_docx_split_points_from_styles(
    blocks: Sequence[Block],
    heading_styles: Sequence[str],
) -> List[Tuple[str, int]]:
    style_set = {s.casefold() for s in heading_styles}
    points = []
    for block in blocks:
        if block.kind != "p" or not block.text.strip() or not block.style_name:
            continue
        if block.style_name.casefold() in style_set:
            points.append((block.text.strip(), block.index))
    return points


def copy_docx_chapter(
    src_docx: Path,
    blocks: Sequence[Block],
    start_idx: int,
    end_idx_exclusive: int,
    out_path: Path,
) -> None:
    shutil.copy2(src_docx, out_path)
    out_doc = Document(str(out_path))
    body = out_doc.element.body
    sectpr = None
    for child in list(body.iterchildren()):
        if child.tag == SECTPR_TAG:
            sectpr = copy.deepcopy(child)
        body.remove(child)
    for block in blocks:
        if start_idx <= block.index < end_idx_exclusive and block.kind != "sectPr":
            body.append(copy.deepcopy(block.xml))
    if sectpr is not None:
        body.append(sectpr)
    out_doc.save(str(out_path))


def infer_output_name(i: int, suffix: str) -> str:
    return f"{i:03d}{suffix}"


def split_docx(
    src_docx: Path,
    output_dir: Path,
    mode: str,
    heading_styles: Sequence[str],
    title_list: Optional[Sequence[str]],
    toc_markers: Sequence[str],
    max_len: int,
) -> List[Path]:
    doc = Document(str(src_docx))
    blocks = docx_blocks(doc)

    if mode == "style":
        points = find_docx_split_points_from_styles(blocks, heading_styles)
    elif mode == "title-list":
        if not title_list:
            die("--title-list is required for mode=title-list")
        points = find_docx_split_points_from_titles(blocks, title_list)
    elif mode == "toc":
        titles, _, toc_end = extract_titles_from_docx_toc(blocks, toc_markers, max_len=max_len)
        if not titles:
            die("Could not extract titles from TOC")
        search_start = toc_end or 0
        points = find_docx_split_points_from_titles(blocks, titles, search_start_index=search_start)
    else:
        die(f"Unsupported mode: {mode}")

    if not points:
        die("No chapter starts were found")

    output_dir.mkdir(parents=True, exist_ok=True)
    created: List[Path] = []
    for i, (_, start_idx) in enumerate(points, 1):
        end_idx = points[i][1] if i < len(points) else max(b.index for b in blocks if b.kind != "sectPr") + 1
        out_path = output_dir / infer_output_name(i, ".docx")
        copy_docx_chapter(src_docx, blocks, start_idx, end_idx, out_path)
        created.append(out_path)
    return created


def pdf_lines(path: Path) -> List[str]:
    text = extract_pdf_text(path)
    return [line.rstrip() for line in text.splitlines()]


def pdf_pages_top_lines(path: Path, max_lines: int = 20) -> List[List[str]]:
    total = pdf_page_count(path)
    pages: List[List[str]] = []
    for page in range(1, total + 1):
        text = extract_pdf_page_text(path, page)
        top = [clean_pdf_line(line) for line in text.splitlines() if clean_pdf_line(line)]
        pages.append(top[:max_lines])
    return pages


def extract_titles_from_pdf_toc(lines: Sequence[str], toc_markers: Sequence[str], max_len: int) -> List[str]:
    markers = {normalize_text(m) for m in toc_markers}
    start = None
    for idx, line in enumerate(lines):
        if normalize_text(line) in markers:
            start = idx + 1
            break
    if start is None:
        return []
    titles: List[str] = []
    for line in lines[start:]:
        title = strip_pdf_toc_entry(line)
        if not title:
            if titles:
                break
            continue
        if is_title_candidate(title, max_len=max_len):
            titles.append(title)
            continue
    return dedupe_preserve_order(merge_title_lines(titles))


def find_pdf_split_points(lines: Sequence[str], titles: Sequence[str]) -> List[Tuple[str, int]]:
    points = []
    pos = 0
    normalized_lines = [normalize_text(line) for line in lines]
    for title in titles:
        target = normalize_text(title)
        found = None
        for idx in range(pos, len(lines)):
            if normalized_lines[idx] == target:
                found = idx
                break
        if found is None:
            continue
        points.append((title, found))
        pos = found + 1
    return points


def find_pdf_chapter_start_pages(
    pages_top_lines: Sequence[Sequence[str]],
    titles: Sequence[str],
    search_start_page: int = 1,
) -> List[Tuple[str, int]]:
    points: List[Tuple[str, int]] = []
    pos = max(0, search_start_page - 1)
    for title in titles:
        target = normalize_text(title)
        found = None
        for page_idx in range(pos, len(pages_top_lines)):
            page_lines = pages_top_lines[page_idx]
            if any(title_matches_page_line(title, line) for line in page_lines):
                found = page_idx + 1  # 1-based pages
                break
        if found is None:
            continue
        points.append((title, found))
        pos = found
    return points


def find_pdf_toc_start_page(
    pages_top_lines: Sequence[Sequence[str]],
    toc_markers: Sequence[str],
) -> Optional[int]:
    markers = {normalize_text(m) for m in toc_markers}
    for page_idx, lines in enumerate(pages_top_lines, 1):
        for line in lines:
            if normalize_text(line) in markers:
                return page_idx
    return None


def extract_pdf_page_range(src_pdf: Path, first_page: int, last_page: int, out_path: Path) -> None:
    subprocess.run(
        [
            "gs",
            "-q",
            "-dNOPAUSE",
            "-dBATCH",
            "-sDEVICE=pdfwrite",
            f"-dFirstPage={first_page}",
            f"-dLastPage={last_page}",
            f"-sOutputFile={out_path}",
            str(src_pdf),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def split_pdf(
    src_pdf: Path,
    output_dir: Path,
    mode: str,
    title_list: Optional[Sequence[str]],
    toc_markers: Sequence[str],
    max_len: int,
) -> List[Path]:
    lines = pdf_lines(src_pdf)
    page_tops = pdf_pages_top_lines(src_pdf)
    if mode == "title-list":
        if not title_list:
            die("--title-list is required for PDF title-list mode")
        titles = list(title_list)
        search_start_page = 1
    elif mode == "toc":
        titles = extract_titles_from_pdf_toc(lines, toc_markers, max_len=max_len)
        if not titles:
            die("Could not extract titles from PDF TOC")
        toc_start_page = find_pdf_toc_start_page(page_tops, toc_markers)
        search_start_page = (toc_start_page + 1) if toc_start_page else 1
    else:
        die("PDF currently supports only mode=toc or mode=title-list")

    points = find_pdf_chapter_start_pages(page_tops, titles, search_start_page=search_start_page)
    if not points:
        die("No chapter starts were found in PDF")

    output_dir.mkdir(parents=True, exist_ok=True)
    created: List[Path] = []
    total_pages = len(page_tops)
    for i, (_, start_page) in enumerate(points, 1):
        end_page = points[i][1] - 1 if i < len(points) else total_pages
        out_path = output_dir / infer_output_name(i, ".pdf")
        extract_pdf_page_range(src_pdf, start_page, end_page, out_path)
        created.append(out_path)
    return created


def resolve_source(path: Path) -> Tuple[Path, Optional[Path]]:
    if not path.exists():
        die(f"Input file not found: {path}")
    if path.suffix.lower() == ".doc":
        converted = convert_doc_to_docx(path)
        return converted, converted.parent
    return path, None


def cleanup_temp(temp_dir: Optional[Path]) -> None:
    if temp_dir and temp_dir.exists():
        shutil.rmtree(temp_dir, ignore_errors=True)


def do_extract_titles(args: argparse.Namespace) -> None:
    src, temp_dir = resolve_source(Path(args.input))
    try:
        if src.suffix.lower() == ".docx":
            doc = Document(str(src))
            blocks = docx_blocks(doc)
            if args.mode == "style":
                titles = extract_titles_from_docx_styles(blocks, args.heading_style)
            elif args.mode == "toc":
                titles, _, _ = extract_titles_from_docx_toc(blocks, args.toc_marker, max_len=args.max_title_length)
            else:
                die("extract-titles supports mode=style or mode=toc for docx/doc")
        elif src.suffix.lower() == ".pdf":
            if args.mode != "toc":
                die("extract-titles for PDF supports only mode=toc")
            titles = extract_titles_from_pdf_toc(pdf_lines(src), args.toc_marker, max_len=args.max_title_length)
        else:
            die(f"Unsupported input type for extract-titles: {src.suffix}")

        if not titles:
            die("No titles were extracted")
        output = Path(args.output) if args.output else None
        text = "\n".join(titles) + "\n"
        if output:
            output.write_text(text, encoding="utf-8")
        else:
            sys.stdout.write(text)
    finally:
        cleanup_temp(temp_dir)


def do_split(args: argparse.Namespace) -> None:
    src, temp_dir = resolve_source(Path(args.input))
    title_list = load_title_list(Path(args.title_list)) if args.title_list else None
    output_dir = Path(args.output_dir)
    try:
        if src.suffix.lower() == ".docx":
            created = split_docx(
                src,
                output_dir=output_dir,
                mode=args.mode,
                heading_styles=args.heading_style,
                title_list=title_list,
                toc_markers=args.toc_marker,
                max_len=args.max_title_length,
            )
        elif src.suffix.lower() == ".pdf":
            created = split_pdf(
                src,
                output_dir=output_dir,
                mode=args.mode,
                title_list=title_list,
                toc_markers=args.toc_marker,
                max_len=args.max_title_length,
            )
        else:
            die(f"Unsupported input type for split: {src.suffix}")

        for path in created:
            print(path)
    finally:
        cleanup_temp(temp_dir)


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Split books into chapter files")
    sub = p.add_subparsers(dest="cmd", required=True)

    for name in ("extract-titles", "split"):
        sp = sub.add_parser(name)
        sp.add_argument("input", help="Path to input file (.doc, .docx, .pdf)")
        sp.add_argument(
            "--mode",
            choices=["style", "title-list", "toc"],
            default="style" if name == "split" else "toc",
            help="How to find chapters",
        )
        sp.add_argument(
            "--heading-style",
            action="append",
            default=list(DEFAULT_HEADING_STYLES),
            help="Paragraph style name that marks H1. Can be repeated.",
        )
        sp.add_argument(
            "--toc-marker",
            action="append",
            default=list(DEFAULT_TOC_MARKERS),
            help="Possible TOC marker lines. Can be repeated.",
        )
        sp.add_argument(
            "--max-title-length",
            type=int,
            default=120,
            help="Max visible length for a TOC title candidate",
        )

    extract_titles = sub.choices["extract-titles"]
    extract_titles.add_argument("--output", help="Write extracted titles to a file")

    split = sub.choices["split"]
    split.add_argument("--title-list", help="One-title-per-line file for mode=title-list")
    split.add_argument("--output-dir", required=True, help="Directory for split chapters")

    return p


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()
    if args.cmd == "extract-titles":
        do_extract_titles(args)
    elif args.cmd == "split":
        do_split(args)
    else:
        parser.error(f"Unknown command: {args.cmd}")


if __name__ == "__main__":
    main()
